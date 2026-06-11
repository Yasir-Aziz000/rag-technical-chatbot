from pathlib import Path
from docx import Document as DocxDocument
from loguru import logger
from src.ingestion.base import BaseDocumentLoader
from src.models.schemas import Document


class DOCXLoader(BaseDocumentLoader):
    def load(self) -> list[Document]:
        doc = DocxDocument(str(self.file_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        if not full_text:
            logger.warning("No text extracted from '{}'", self.source_name)
            return []

        # DOCX has no native page numbers — return as single document
        documents = [
            Document(
                content=full_text,
                source=self.source_name,
                page=None,
                doc_type="docx",
            )
        ]

        logger.info("Loaded 1 document from '{}'", self.source_name)
        return documents
